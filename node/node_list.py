
from state.state_list import DocState
from struct_output.output_list import (
    ParsedDocWithMetadata, 
    DocSectionWithMetadata, DocSectionMetadata, TableData, FunctionSection
)
from langchain_core.documents import Document as LCDocument

from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from typing import List, Dict, Optional


class WordDocumentParser:
    def __init__(self, doc_path: str):
        self.doc = DocxDocument(doc_path)
        
    def parse_section_3(self) -> ParsedDocWithMetadata:
        in_section_3 = False
        current_h2 = None
        current_h3 = None
        current_h4 = None
        current_content = []
        current_tables = []
        current_function_sections = []
        current_section_type = None
        section_type_content = []

        result = []

        def flush_section():
            nonlocal current_h2, current_h3, current_h4, current_content, current_tables, current_function_sections, current_section_type, section_type_content
            if not current_h3:
                return
            if section_type_content and current_section_type:
                current_function_sections.append(
                    FunctionSection(
                        section_type=current_section_type,
                        content='\n'.join(section_type_content),
                        tables=[]
                    )
                )
            section_data = self._create_section_data(
                current_h2, current_h3, current_h4,
                current_content, current_function_sections, current_tables
            )
            result.append(section_data)

        def reset_section():
            nonlocal current_h4, current_content, current_tables, current_function_sections, current_section_type, section_type_content
            current_h4 = None
            current_content = []
            current_tables = []
            current_function_sections = []
            current_section_type = None
            section_type_content = []

        for child in self.doc.element.body:
            tag = child.tag

            if tag == qn('w:p'):
                para = Paragraph(child, self.doc)
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name

                if style_name == 'Heading 1':
                    if '功能分析' in text:
                        in_section_3 = True
                        continue
                    elif in_section_3:
                        flush_section()
                        break

                if not in_section_3:
                    continue

                if style_name == 'Heading 2':
                    flush_section()
                    current_h2 = text
                    reset_section()

                elif style_name == 'Heading 3':
                    flush_section()
                    current_h3 = text
                    reset_section()

                elif style_name == 'Heading 4':
                    if section_type_content and current_section_type:
                        current_function_sections.append(
                            FunctionSection(
                                section_type=current_section_type,
                                content='\n'.join(section_type_content),
                                tables=[]
                            )
                        )
                    current_h4 = text
                    current_section_type = text
                    section_type_content = []

                elif style_name == 'Heading 5':
                    if section_type_content and current_section_type:
                        current_function_sections.append(
                            FunctionSection(
                                section_type=current_section_type,
                                content='\n'.join(section_type_content),
                                tables=[]
                            )
                        )
                    current_section_type = text
                    section_type_content = []

                elif style_name.startswith('Heading'):
                    pass

                else:
                    if text:
                        current_content.append(text)
                        if current_section_type:
                            section_type_content.append(text)

            elif tag == qn('w:tbl'):
                if not in_section_3:
                    continue
                table = Table(child, self.doc)
                table_data = self._extract_table(table)
                if table_data:
                    current_tables.append(table_data)

        flush_section()

        return ParsedDocWithMetadata(
            sections=result,
            total_count=len(result)
        )
    
    def _extract_table(self, table) -> Optional[TableData]:
        try:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if not rows:
                return None

            header_start = 0
            while header_start < len(rows):
                row = rows[header_start]
                if len(set(row)) == 1 and len(row) > 1:
                    header_start += 1
                else:
                    break

            if header_start >= len(rows):
                return None

            headers = rows[header_start]
            data_rows = rows[header_start + 1:]

            if headers and data_rows:
                return TableData(headers=headers, rows=data_rows)
            return None
        except Exception as e:
            print(f"表格提取失败: {e}")
            return None
    
    def _create_section_data(
        self, 
        h2: str, 
        h3: str, 
        h4: Optional[str], 
        content: List[str],
        function_sections: List[FunctionSection],
        tables: List[TableData]
    ) -> DocSectionWithMetadata:
        content_text = '\n'.join(content)
        
        metadata = DocSectionMetadata(
            level_1="功能分析",
            level_2=h2,
            level_3=h3,
            level_4=h4
        )
        
        title_parts = [h3]
        if h4:
            title_parts.append(h4)
        full_title = ' - '.join(title_parts)
        
        return DocSectionWithMetadata(
            title=full_title,
            level=3,
            content=content_text,
            metadata=metadata,
            function_sections=function_sections,
            tables=tables
        )


def word_parser_node(state: DocState) -> Dict:
    file_path = state["file_path"]
    parser = WordDocumentParser(file_path)
    parsed_data = parser.parse_section_3()
    
    return {"parsed_data": parsed_data}

def word_indexer_node(state: DocState) -> Dict:
    sections = state["parsed_data"].sections
    documents = []
    
    for sec in sections:
        combined_text = f"标题：{sec.title}\n层级路径：{sec.metadata.level_1} > {sec.metadata.level_2} > {sec.metadata.level_3}"
        if sec.metadata.level_4:
            combined_text += f" > {sec.metadata.level_4}"
        combined_text += f"\n\n内容：{sec.content}"
        
        if sec.function_sections:
            combined_text += "\n\n功能分解："
            for fs in sec.function_sections:
                combined_text += f"\n【{fs.section_type}】\n{fs.content}"
        
        if sec.tables:
            combined_text += "\n\n表格数据："
            for i, table in enumerate(sec.tables, 1):
                combined_text += f"\n表格{i}: {', '.join(table.headers)}"
                for row in table.rows[:2]:
                    combined_text += f"\n  {', '.join(row)}"
        
        metadata = {
            "title": sec.title,
            "level": sec.level,
            "level_1": sec.metadata.level_1,
            "level_2": sec.metadata.level_2,
            "level_3": sec.metadata.level_3,
            "level_4": sec.metadata.level_4 or "",
            "function_section_types": ", ".join([fs.section_type for fs in sec.function_sections]) if sec.function_sections else "",
            "table_count": str(len(sec.tables))
        }
        
        documents.append(LCDocument(page_content=combined_text, metadata=metadata))
    
    db = Chroma.from_documents(
        documents,
        HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"),
        persist_directory="./chroma_db",
        collection_name="word_index"
    )
    
    return {"index_status": "Completed"}